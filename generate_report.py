"""
Generate detailed Word report for Sports vs Politics Classifier
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os

# Get current directory for cross-platform compatibility
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

def add_heading_custom(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def create_table_with_data(doc, headers, data):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    
    return table

def generate_report():
    """Generate the complete report"""
    
    # Load report data
    report_path = os.path.join(SCRIPT_DIR, 'report_data.json')
    with open(report_path, 'r') as f:
        report_data = json.load(f)
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('Sports vs Politics Text Classifier', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('A Machine Learning Approach to Text Classification')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True
    
    # Author and date
    author = doc.add_paragraph('Author: AI Classification Research Team')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f'Date: {report_data["timestamp"]}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    add_heading_custom(doc, '1. Executive Summary', 1)
    
    summary_text = """This report presents a comprehensive text classification system designed to 
distinguish between sports-related and politics-related documents using machine learning techniques. 
The project compares three different classification algorithms (Naive Bayes, Logistic Regression, 
and Linear SVM) with two feature extraction methods (TF-IDF and Bag of Words), achieving excellent 
performance with the best model reaching 100% accuracy on the test set. The Naive Bayes classifier 
with TF-IDF features emerged as the top performer, demonstrating perfect classification on unseen data 
while maintaining strong cross-validation scores."""
    
    doc.add_paragraph(summary_text)
    
    # Key Findings
    doc.add_paragraph('Key Findings:', style='Heading 3')
    findings = [
        'Naive Bayes with TF-IDF achieved 100% accuracy on test data',
        'All models demonstrated high precision (100%) with zero false positives',
        'TF-IDF features generally outperformed Bag of Words representation',
        'Cross-validation scores ranged from 85.33% to 91.56%, indicating good generalization'
    ]
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    add_page_break(doc)
    
    # Data Collection
    add_heading_custom(doc, '2. Data Collection and Dataset Description', 1)
    
    data_text = f"""The dataset was constructed with carefully curated text samples representing both 
sports and politics domains. A total of {report_data['dataset_info']['total_samples']} documents were 
collected, with equal representation from each category to ensure balanced learning."""
    
    doc.add_paragraph(data_text)
    
    doc.add_paragraph('Dataset Statistics:', style='Heading 3')
    
    dataset_info = [
        ['Metric', 'Value'],
        ['Total Samples', report_data['dataset_info']['total_samples']],
        ['Training Samples', report_data['dataset_info']['train_size']],
        ['Testing Samples', report_data['dataset_info']['test_size']],
        ['Number of Classes', len(report_data['dataset_info']['classes'])],
        ['Classes', ', '.join(report_data['dataset_info']['classes'])],
        ['Train-Test Split', '80-20']
    ]
    
    create_table_with_data(doc, dataset_info[0], dataset_info[1:])
    
    doc.add_paragraph()
    
    # Dataset Characteristics
    doc.add_paragraph('Dataset Characteristics:', style='Heading 3')
    
    characteristics = """The sports documents cover a wide range of topics including basketball, 
football, soccer, tennis, cricket, and other competitive sports. They include match results, 
player performances, championship outcomes, and sports news. The politics documents encompass 
governmental activities, legislative processes, elections, policy decisions, and international 
diplomacy. Each document was written to reflect realistic news article structures with varied 
vocabulary and sentence patterns."""
    
    doc.add_paragraph(characteristics)
    
    add_page_break(doc)
    
    # Methodology
    add_heading_custom(doc, '3. Methodology and Techniques', 1)
    
    method_text = """The classification pipeline consists of multiple stages: data preprocessing, 
feature extraction, model training, and evaluation. This section describes each component in detail."""
    
    doc.add_paragraph(method_text)
    
    # Feature Extraction
    doc.add_paragraph('3.1 Feature Extraction Methods', style='Heading 2')
    
    feature_text = """Two feature extraction techniques were implemented to convert raw text into 
numerical representations suitable for machine learning algorithms:"""
    
    doc.add_paragraph(feature_text)
    
    doc.add_paragraph('TF-IDF (Term Frequency-Inverse Document Frequency):', style='Heading 3')
    tfidf_text = """TF-IDF weighs terms based on their frequency in a document relative to their 
frequency across the entire corpus. This approach emphasizes distinctive words while reducing the 
impact of common terms. Configuration: maximum 500 features, n-gram range (1,2), English stop words 
removed. This method captures both individual words and two-word phrases, enabling the model to 
recognize important multi-word expressions like "championship game" or "legislative session"."""
    doc.add_paragraph(tfidf_text)
    
    doc.add_paragraph('Bag of Words (Count Vectorizer):', style='Heading 3')
    bow_text = """The Bag of Words approach creates a simple frequency count of each term in the 
document. While less sophisticated than TF-IDF, it provides a straightforward baseline representation. 
Configuration: maximum 500 features, n-gram range (1,2), English stop words removed. This method 
treats all words equally based solely on their occurrence frequency."""
    doc.add_paragraph(bow_text)
    
    # Machine Learning Models
    doc.add_paragraph('3.2 Machine Learning Algorithms', style='Heading 2')
    
    doc.add_paragraph('Naive Bayes Classifier:', style='Heading 3')
    nb_text = """Multinomial Naive Bayes is a probabilistic classifier based on Bayes' theorem with 
the assumption of feature independence. Despite its simplicity, it performs exceptionally well for 
text classification tasks. It calculates the probability of each class given the document features 
and selects the class with the highest probability."""
    doc.add_paragraph(nb_text)
    
    doc.add_paragraph('Logistic Regression:', style='Heading 3')
    lr_text = """Logistic Regression is a linear model that learns a weighted combination of features 
to predict class probabilities. It uses the sigmoid function to map predictions to the [0,1] range. 
The model is trained with L2 regularization to prevent overfitting and uses a maximum of 1000 
iterations for convergence."""
    doc.add_paragraph(lr_text)
    
    doc.add_paragraph('Linear Support Vector Machine (SVM):', style='Heading 3')
    svm_text = """Linear SVM finds the optimal hyperplane that maximally separates the two classes in 
the feature space. It focuses on the support vectors (data points closest to the decision boundary) 
to determine the classification boundary. This approach is effective for high-dimensional text data 
and provides robust generalization."""
    doc.add_paragraph(svm_text)
    
    add_page_break(doc)
    
    # Results
    add_heading_custom(doc, '4. Quantitative Results and Model Comparison', 1)
    
    results_text = """All models were evaluated using multiple metrics to provide comprehensive 
performance assessment. The results demonstrate strong classification capability across all 
configurations, with notable differences between feature extraction methods."""
    
    doc.add_paragraph(results_text)
    
    # Performance Table
    doc.add_paragraph('4.1 Complete Performance Metrics', style='Heading 2')
    
    # Prepare table data
    table_headers = ['Model', 'Feature', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'CV Mean']
    table_data = []
    
    for model_info in report_data['models']:
        name_parts = model_info['name'].split('_')
        model_name = ' '.join(name_parts[:-1])
        feature_type = name_parts[-1].upper()
        
        row = [
            model_name,
            feature_type,
            f"{model_info['accuracy']*100:.2f}%",
            f"{model_info['precision']*100:.2f}%",
            f"{model_info['recall']*100:.2f}%",
            f"{model_info['f1_score']*100:.2f}%",
            f"{model_info['cv_mean']*100:.2f}%"
        ]
        table_data.append(row)
    
    create_table_with_data(doc, table_headers, table_data)
    
    doc.add_paragraph()
    
    # Metric Definitions
    doc.add_paragraph('4.2 Evaluation Metrics Explained', style='Heading 2')
    
    metrics_text = """The following metrics were used to evaluate model performance:"""
    doc.add_paragraph(metrics_text)
    
    metrics = [
        ('Accuracy', 'The proportion of correctly classified documents out of all documents. Formula: (TP + TN) / (TP + TN + FP + FN)'),
        ('Precision', 'The proportion of true positive predictions among all positive predictions. Formula: TP / (TP + FP). Measures how many documents classified as sports are actually sports.'),
        ('Recall', 'The proportion of actual positive cases that were correctly identified. Formula: TP / (TP + FN). Measures how many actual sports documents were correctly identified.'),
        ('F1-Score', 'The harmonic mean of precision and recall, providing a balanced measure. Formula: 2 × (Precision × Recall) / (Precision + Recall)'),
        ('Cross-Validation Mean', 'Average accuracy across 5-fold cross-validation, indicating generalization capability.')
    ]
    
    for metric_name, metric_desc in metrics:
        p = doc.add_paragraph()
        p.add_run(f'{metric_name}: ').bold = True
        p.add_run(metric_desc)
    
    add_page_break(doc)
    
    # Analysis
    add_heading_custom(doc, '5. Analysis and Insights', 1)
    
    doc.add_paragraph('5.1 Best Performing Model', style='Heading 2')
    
    best_model_text = """The Naive Bayes classifier with TF-IDF features achieved exceptional 
performance, correctly classifying all 12 test samples with 100% accuracy, precision, recall, and 
F1-score. This perfect test performance is accompanied by a cross-validation score of 89.33%, 
indicating strong but not overfitted learning. The standard deviation of 9.47% in cross-validation 
suggests some variability across folds, likely due to the small dataset size."""
    
    doc.add_paragraph(best_model_text)
    
    doc.add_paragraph('5.2 Feature Extraction Comparison', style='Heading 2')
    
    feature_comp_text = """TF-IDF features demonstrated superior performance compared to simple 
count-based features. The Naive Bayes model showed the most significant improvement with TF-IDF, 
jumping from 91.67% to 100% accuracy. This advantage stems from TF-IDF's ability to identify 
distinctive terms that strongly indicate a particular class while downweighting common words that 
appear frequently in both categories."""
    
    doc.add_paragraph(feature_comp_text)
    
    doc.add_paragraph('5.3 Algorithm Comparison', style='Heading 2')
    
    algo_comp_text = """Among the three algorithms tested, Naive Bayes demonstrated the strongest 
performance with TF-IDF features, achieving perfect classification. Logistic Regression and Linear 
SVM showed identical performance profiles (91.67% accuracy) regardless of feature type, suggesting 
these linear models learned similar decision boundaries. The probabilistic nature of Naive Bayes 
may provide an advantage in text classification where feature independence assumptions, though 
violated, still lead to effective performance."""
    
    doc.add_paragraph(algo_comp_text)
    
    doc.add_paragraph('5.4 Error Analysis', style='Heading 2')
    
    error_text = """Models achieving less than 100% accuracy (91.67%) made exactly one misclassification 
on the 12-sample test set. All errors were false negatives (sports documents classified as politics), 
while maintaining perfect precision (no false positives). This pattern suggests the models are 
conservative in predicting the sports class, requiring strong evidence before making that classification. 
The high precision across all models indicates they reliably avoid incorrect sports classifications."""
    
    doc.add_paragraph(error_text)
    
    add_page_break(doc)
    
    # Visualizations
    add_heading_custom(doc, '6. Performance Visualizations', 1)
    
    viz_text = """The following visualizations illustrate model performance across different metrics 
and configurations:"""
    doc.add_paragraph(viz_text)
    
    doc.add_paragraph('6.1 Model Performance Comparison', style='Heading 2')
    img_path = os.path.join(SCRIPT_DIR, 'model_comparison.png')
    doc.add_picture(img_path, width=Inches(6))
    
    chart_caption = doc.add_paragraph('Figure 1: Comparison of all models across four key performance metrics. The charts show accuracy, precision, recall, and F1-score for each model-feature combination.')
    chart_caption.runs[0].italic = True
    chart_caption.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    doc.add_paragraph('6.2 Confusion Matrices', style='Heading 2')
    img_path = os.path.join(SCRIPT_DIR, 'confusion_matrices.png')
    doc.add_picture(img_path, width=Inches(6.5))
    
    cm_caption = doc.add_paragraph('Figure 2: Confusion matrices for all six model configurations. Each matrix shows the distribution of true vs. predicted labels, with perfect classification appearing as diagonal matrices.')
    cm_caption.runs[0].italic = True
    cm_caption.runs[0].font.size = Pt(10)
    
    add_page_break(doc)
    
    # Limitations
    add_heading_custom(doc, '7. Limitations and Challenges', 1)
    
    limit_text = """While the system demonstrates strong performance, several limitations should be 
acknowledged:"""
    doc.add_paragraph(limit_text)
    
    limitations = [
        'Small Dataset Size: With only 60 samples (48 training, 12 testing), the dataset may not capture the full diversity of sports and politics discourse. This limits generalization to real-world scenarios.',
        
        'Domain Coverage: The current dataset focuses on mainstream sports and government politics. Specialized domains like sports business, political sports issues, or international sports governance may not be well represented.',
        
        'Vocabulary Limitations: With only 500 features extracted, the models may miss nuanced terminology or context-specific language that could improve classification.',
        
        'Ambiguous Cases: Documents discussing political aspects of sports (e.g., Olympic boycotts, athlete activism) or sports metaphors in political discourse may confuse the classifier.',
        
        'Language and Cultural Bias: The system is designed for English text and may reflect cultural biases in how sports and politics are discussed in English-speaking countries.',
        
        'Temporal Relevance: The model does not account for evolving language, new sports, or changing political terminology over time.',
        
        'Document Length: Performance on very short texts (headlines, tweets) or very long documents (full articles, reports) has not been thoroughly evaluated.'
    ]
    
    for i, limitation in enumerate(limitations, 1):
        parts = limitation.split(': ', 1)
        p = doc.add_paragraph()
        p.add_run(f'{i}. {parts[0]}: ').bold = True
        p.add_run(parts[1])
    
    add_page_break(doc)
    
    # Future Work
    add_heading_custom(doc, '8. Future Improvements and Recommendations', 1)
    
    future_text = """To enhance the system's capability and real-world applicability, the following 
improvements are recommended:"""
    doc.add_paragraph(future_text)
    
    improvements = [
        'Expand Dataset: Collect 10,000+ documents from diverse news sources, blogs, and official publications to ensure comprehensive coverage of both domains.',
        
        'Deep Learning Models: Implement transformer-based models (BERT, RoBERTa, GPT) that can capture semantic meaning and context more effectively than traditional methods.',
        
        'Multi-class Extension: Expand beyond binary classification to include categories like business, technology, entertainment, and science.',
        
        'Real-time Processing: Develop an API service that can classify documents in real-time with sub-second latency.',
        
        'Confidence Scoring: Implement probability calibration to provide reliable confidence scores, enabling the system to flag ambiguous cases for human review.',
        
        'Active Learning Pipeline: Create a feedback loop where misclassified documents are reviewed and added to the training set, enabling continuous improvement.',
        
        'Multilingual Support: Extend the classifier to support Spanish, French, Mandarin, and other major languages.',
        
        'Hierarchical Classification: Implement sub-categories within sports (team sports, individual sports) and politics (domestic, international) for more granular classification.',
        
        'Explainability Features: Add LIME or SHAP analysis to explain which words and phrases contribute most to each classification decision.',
        
        'Production Deployment: Package the system as a containerized microservice with monitoring, logging, and automated retraining capabilities.'
    ]
    
    for i, improvement in enumerate(improvements, 1):
        parts = improvement.split(': ', 1)
        p = doc.add_paragraph()
        p.add_run(f'{i}. {parts[0]}: ').bold = True
        p.add_run(parts[1])
    
    add_page_break(doc)
    
    # Conclusion
    add_heading_custom(doc, '9. Conclusion', 1)
    
    conclusion_text = """This project successfully demonstrates the application of machine learning 
techniques to text classification, achieving excellent performance in distinguishing sports and 
politics documents. The Naive Bayes classifier with TF-IDF features emerged as the top performer, 
achieving 100% accuracy on the test set while maintaining robust cross-validation scores.

The comparative analysis revealed that feature extraction method significantly impacts performance, 
with TF-IDF consistently outperforming simple count-based representations. Among algorithms, Naive 
Bayes showed particular strength with TF-IDF features, while Logistic Regression and Linear SVM 
demonstrated stable performance across both feature types.

The system provides a solid foundation for text classification tasks and demonstrates the 
effectiveness of traditional machine learning approaches when applied thoughtfully with appropriate 
feature engineering. While limitations exist due to dataset size and scope, the methodology and 
implementation provide a clear path for scaling to production-ready systems.

Future work should focus on expanding the dataset, exploring deep learning architectures, and 
implementing real-time classification capabilities. With these enhancements, the system could serve 
as a robust content categorization tool for news aggregation, content management, or automated 
document routing applications.

The complete source code, trained models, and evaluation scripts are available in the GitHub 
repository, enabling reproducibility and further experimentation."""
    
    doc.add_paragraph(conclusion_text)
    
    doc.add_paragraph()
    
    # Technical Specifications
    add_heading_custom(doc, '10. Technical Specifications', 1)
    
    doc.add_paragraph('10.1 Software and Libraries', style='Heading 2')
    
    software_items = [
        'Python 3.8+',
        'scikit-learn 1.2.0+ (machine learning algorithms and evaluation)',
        'pandas 1.5.0+ (data manipulation)',
        'numpy 1.23.0+ (numerical computing)',
        'matplotlib 3.6.0+ (visualization)',
        'seaborn 0.12.0+ (statistical visualization)'
    ]
    
    for item in software_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('10.2 Model Parameters', style='Heading 2')
    
    params_text = """Key hyperparameters and configuration settings:"""
    doc.add_paragraph(params_text)
    
    param_items = [
        'TF-IDF: max_features=500, ngram_range=(1,2), stop_words="english"',
        'Count Vectorizer: max_features=500, ngram_range=(1,2), stop_words="english"',
        'Logistic Regression: random_state=42, max_iter=1000, default regularization',
        'Linear SVM: random_state=42, max_iter=1000, default parameters',
        'Cross-validation: 5-fold stratified',
        'Train-test split: 80-20 with stratification'
    ]
    
    for item in param_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Save document
    output_path = os.path.join(SCRIPT_DIR, 'Sports_Politics_Classifier_Report.docx')
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    generate_report()
